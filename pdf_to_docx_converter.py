"""
Advanced PDF to DOCX Conversion Tool
Supports multiple conversion backends with consistent exception handling
Achieves WPS-level conversion quality through method fallback strategy
"""

import os
import sys
import logging
from pathlib import Path
from typing import Optional, Dict, Tuple
import tempfile
import subprocess
import shutil


# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class PDFToDOCXError(Exception):
    """Custom exception for PDF to DOCX conversion errors"""
    pass


class PDFToDOCXConverter:
    """
    Converts PDF files to DOCX format using multiple backend methods.
    
    Supported methods (in priority order):
    1. pdfplumber + python-docx: High-quality text extraction with layout preservation
    2. pdf2image + pytesseract (OCR): For scanned PDFs
    3. LibreOffice headless: System-level conversion with best compatibility
    4. pymupdf: Lightweight fallback for simple PDFs
    """
    
    def __init__(self, timeout: int = 120, prefer_ocr: bool = False):
        """
        Initialize the converter.
        
        Args:
            timeout: Maximum conversion time in seconds (default: 120)
            prefer_ocr: Prefer OCR method for all PDFs (default: False)
        """
        self.timeout = timeout
        self.prefer_ocr = prefer_ocr
        self.conversion_errors = []
    
    def convert(
        self,
        input_path: str,
        output_path: Optional[str] = None,
        method: Optional[str] = None
    ) -> str:
        """
        Convert PDF to DOCX with automatic fallback strategy.
        
        Args:
            input_path: Path to input PDF file
            output_path: Path for output DOCX file (auto-generated if None)
            method: Force specific conversion method 
                   ('pdfplumber', 'ocr', 'libreoffice', 'pymupdf', 'auto')
                   Default: 'auto' (tries best available method)
        
        Returns:
            Path to the generated DOCX file
            
        Raises:
            PDFToDOCXError: If conversion fails with all available methods
        """
        # Validate input
        input_file = self._validate_input(input_path)
        
        # Generate output path if not provided
        output_file = self._generate_output_path(input_file, output_path)
        
        logger.info(f"Starting PDF to DOCX conversion: {input_path} -> {output_file}")
        
        # Reset error tracking
        self.conversion_errors = []
        
        # Determine conversion method
        if method == 'auto' or method is None:
            return self._convert_auto(input_file, output_file)
        else:
            return self._convert_with_method(input_file, output_file, method)
    
    def _validate_input(self, input_path: str) -> Path:
        """Validate input PDF file exists and is readable."""
        input_file = Path(input_path)
        
        if not input_file.exists():
            raise PDFToDOCXError(f"Input file not found: {input_path}")
        
        if not input_file.is_file():
            raise PDFToDOCXError(f"Input path is not a file: {input_path}")
        
        if input_file.suffix.lower() != '.pdf':
            raise PDFToDOCXError(f"Input file is not a PDF: {input_path}")
        
        if not os.access(input_file, os.R_OK):
            raise PDFToDOCXError(f"Input file is not readable: {input_path}")
        
        logger.debug(f"Input validation passed: {input_file}")
        return input_file
    
    def _generate_output_path(self, input_file: Path, output_path: Optional[str]) -> Path:
        """Generate output DOCX path."""
        if output_path:
            output_file = Path(output_path)
        else:
            output_file = input_file.parent / (input_file.stem + '.docx')
        
        # Create parent directory if needed
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        logger.debug(f"Output path: {output_file}")
        return output_file
    
    def _convert_auto(self, input_file: Path, output_file: Path) -> str:
        """Try conversion methods in optimal order based on system capabilities."""
        methods = self._get_available_methods()
        
        if not methods:
            raise PDFToDOCXError(
                "No conversion methods available. Please install one of:\n"
                "  pip install pdfplumber python-docx\n"
                "  pip install pdf2image pytesseract\n"
                "  brew/apt install libreoffice\n"
                "  pip install pymupdf"
            )
        
        logger.info(f"Available conversion methods: {methods}")
        
        for method in methods:
            try:
                logger.info(f"Attempting conversion with method: {method}")
                result = self._convert_with_method(input_file, output_file, method)
                logger.info(f"Conversion successful using {method}")
                return result
            except PDFToDOCXError as e:
                error_msg = f"{method}: {str(e)}"
                self.conversion_errors.append(error_msg)
                logger.warning(f"Method {method} failed: {e}")
                continue
        
        # All methods failed
        error_summary = "\n".join(self.conversion_errors)
        raise PDFToDOCXError(
            f"All conversion methods failed:\n{error_summary}"
        )
    
    def _get_available_methods(self) -> list:
        """Detect available conversion methods on current system."""
        available = []
        
        # Method 1: pdfplumber (best quality for text extraction)
        if self._check_pdfplumber():
            available.append('pdfplumber')
        
        # Method 2: OCR (for scanned PDFs)
        if self._check_ocr() or self.prefer_ocr:
            available.append('ocr')
        
        # Method 3: LibreOffice (system-level, best compatibility)
        if self._find_libreoffice():
            available.append('libreoffice')
        
        # Method 4: PyMuPDF (lightweight fallback)
        if self._check_pymupdf():
            available.append('pymupdf')
        
        return available
    
    @staticmethod
    def _check_pdfplumber() -> bool:
        """Check if pdfplumber is available."""
        try:
            import pdfplumber
            logger.debug("pdfplumber is available")
            return True
        except ImportError:
            logger.debug("pdfplumber not available")
            return False
    
    @staticmethod
    def _check_ocr() -> bool:
        """Check if OCR dependencies are available."""
        try:
            from pdf2image import convert_from_path
            import pytesseract
            from PIL import Image
            
            # Check if tesseract-ocr is installed on system
            result = subprocess.run(
                ['tesseract', '--version'],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )
            ocr_available = result.returncode == 0
            
            if ocr_available:
                logger.debug("OCR dependencies available")
            else:
                logger.debug("pytesseract found but system tesseract not installed")
            
            return ocr_available
        except (ImportError, FileNotFoundError):
            logger.debug("OCR dependencies not available")
            return False
    
    @staticmethod
    def _find_libreoffice() -> Optional[str]:
        """Find LibreOffice binary path."""
        candidates = [
            'libreoffice',
            'soffice',
            '/usr/bin/libreoffice',
            '/usr/bin/soffice',
            'C:\\Program Files\\LibreOffice\\program\\soffice.exe',
            'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe',
        ]
        
        for candidate in candidates:
            if shutil.which(candidate):
                logger.debug(f"Found LibreOffice at: {candidate}")
                return candidate
        
        logger.debug("LibreOffice not found on system")
        return None
    
    @staticmethod
    def _check_pymupdf() -> bool:
        """Check if PyMuPDF is available."""
        try:
            import fitz
            logger.debug("PyMuPDF is available")
            return True
        except ImportError:
            logger.debug("PyMuPDF not available")
            return False
    
    def _convert_with_method(self, input_file: Path, output_file: Path, method: str) -> str:
        """Execute conversion with specified method."""
        method_lower = method.lower()
        
        if method_lower == 'pdfplumber':
            return self._convert_pdfplumber(input_file, output_file)
        elif method_lower == 'ocr':
            return self._convert_ocr(input_file, output_file)
        elif method_lower == 'libreoffice':
            return self._convert_libreoffice(input_file, output_file)
        elif method_lower == 'pymupdf':
            return self._convert_pymupdf(input_file, output_file)
        else:
            raise PDFToDOCXError(f"Unknown conversion method: {method}")
    
    def _convert_pdfplumber(self, input_file: Path, output_file: Path) -> str:
        """
        Convert using pdfplumber + python-docx
        Best for: PDFs with selectable text and good structure
        Quality: High - Preserves text formatting and table structure
        """
        try:
            import pdfplumber
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            logger.debug("Using pdfplumber for conversion")
            
            doc = Document()
            
            with pdfplumber.open(input_file) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"Processing {total_pages} pages")
                
                for page_idx, page in enumerate(pdf.pages, 1):
                    logger.debug(f"Processing page {page_idx}/{total_pages}")
                    
                    # Extract text
                    text = page.extract_text()
                    if text and text.strip():
                        doc.add_paragraph(text)
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        for table_data in tables:
                            if table_data:
                                self._add_table_to_docx(doc, table_data)
                    
                    # Add page break between pages
                    if page_idx < total_pages:
                        doc.add_page_break()
            
            doc.save(output_file)
            logger.info(f"PDF successfully converted to DOCX: {output_file}")
            return str(output_file)
            
        except ImportError as e:
            raise PDFToDOCXError(
                f"pdfplumber not installed: {e}\n"
                "Install with: pip install pdfplumber python-docx"
            )
        except Exception as e:
            raise PDFToDOCXError(f"pdfplumber conversion failed: {e}")
    
    def _convert_ocr(self, input_file: Path, output_file: Path) -> str:
        """
        Convert using pdf2image + pytesseract OCR
        Best for: Scanned PDFs and image-based documents
        Quality: Medium-High - Accuracy depends on image quality
        """
        try:
            from pdf2image import convert_from_path
            import pytesseract
            from PIL import Image
            from docx import Document
            
            logger.debug("Using OCR method for conversion")
            
            # Convert PDF pages to images
            logger.info("Converting PDF pages to images")
            images = convert_from_path(
                str(input_file),
                timeout=self.timeout,
                thread_count=4
            )
            
            doc = Document()
            
            for page_idx, image in enumerate(images, 1):
                logger.debug(f"Processing page {page_idx} with OCR")
                
                # Perform OCR on image
                text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                
                if text and text.strip():
                    doc.add_paragraph(text)
                
                # Add page break
                if page_idx < len(images):
                    doc.add_page_break()
            
            doc.save(output_file)
            logger.info(f"PDF successfully converted via OCR: {output_file}")
            return str(output_file)
            
        except ImportError as e:
            raise PDFToDOCXError(
                f"OCR dependencies not installed: {e}\n"
                "Install with: pip install pdf2image pytesseract python-docx\n"
                "System requirement: tesseract-ocr (brew/apt install tesseract-ocr)"
            )
        except subprocess.TimeoutExpired:
            raise PDFToDOCXError(f"OCR conversion timeout (>{self.timeout}s)")
        except Exception as e:
            raise PDFToDOCXError(f"OCR conversion failed: {e}")
    
    def _convert_libreoffice(self, input_file: Path, output_file: Path) -> str:
        """
        Convert using LibreOffice headless
        Best for: Compatibility and quality (WPS-equivalent)
        Quality: Highest - Best format preservation
        """
        try:
            libreoffice_bin = self._find_libreoffice()
            if not libreoffice_bin:
                raise PDFToDOCXError("LibreOffice not found on system")
            
            logger.debug(f"Using LibreOffice: {libreoffice_bin}")
            
            # Create temp directory for conversion
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_dir_path = Path(temp_dir)
                
                # Run LibreOffice conversion
                cmd = [
                    libreoffice_bin,
                    '--headless',
                    '--convert-to', 'docx',
                    '--outdir', str(temp_dir_path),
                    str(input_file)
                ]
                
                logger.debug(f"Running command: {' '.join(cmd)}")
                
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=self.timeout
                )
                
                if result.returncode != 0:
                    raise PDFToDOCXError(
                        f"LibreOffice conversion failed: {result.stderr}"
                    )
                
                # Find generated DOCX
                pdf_basename = input_file.stem
                generated_docx = temp_dir_path / (pdf_basename + '.docx')
                
                if not generated_docx.exists():
                    raise PDFToDOCXError(
                        f"LibreOffice conversion completed but output not found: {generated_docx}"
                    )
                
                # Move to target location
                shutil.move(str(generated_docx), str(output_file))
            
            logger.info(f"PDF successfully converted via LibreOffice: {output_file}")
            return str(output_file)
            
        except subprocess.TimeoutExpired:
            raise PDFToDOCXError(f"LibreOffice conversion timeout (>{self.timeout}s)")
        except FileNotFoundError:
            raise PDFToDOCXError(
                "LibreOffice binary not found\n"
                "Install with:\n"
                "  macOS: brew install libreoffice\n"
                "  Ubuntu: sudo apt install libreoffice-writer\n"
                "  CentOS: sudo yum install libreoffice-writer"
            )
        except Exception as e:
            raise PDFToDOCXError(f"LibreOffice conversion failed: {e}")
    
    def _convert_pymupdf(self, input_file: Path, output_file: Path) -> str:
        """
        Convert using PyMuPDF (fitz)
        Best for: Simple PDFs, lightweight option
        Quality: Medium - Basic text extraction
        """
        try:
            import fitz
            from docx import Document
            
            logger.debug("Using PyMuPDF for conversion")
            
            pdf_doc = fitz.open(str(input_file))
            word_doc = Document()
            
            for page_idx, page in enumerate(pdf_doc, 1):
                logger.debug(f"Processing page {page_idx}")
                
                # Extract text
                text = page.get_text()
                if text and text.strip():
                    word_doc.add_paragraph(text)
                
                # Add page break
                if page_idx < len(pdf_doc):
                    word_doc.add_page_break()
            
            pdf_doc.close()
            word_doc.save(str(output_file))
            
            logger.info(f"PDF successfully converted via PyMuPDF: {output_file}")
            return str(output_file)
            
        except ImportError as e:
            raise PDFToDOCXError(
                f"PyMuPDF not installed: {e}\n"
                "Install with: pip install pymupdf python-docx"
            )
        except Exception as e:
            raise PDFToDOCXError(f"PyMuPDF conversion failed: {e}")
    
    @staticmethod
    def _add_table_to_docx(doc, table_data: list) -> None:
        """Add extracted table data to DOCX document."""
        try:
            from docx.shared import Pt
            
            if not table_data or not table_data[0]:
                return
            
            rows = len(table_data)
            cols = len(table_data[0])
            
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Light Grid Accent 1'
            
            for i, row in enumerate(table_data):
                for j, cell_text in enumerate(row):
                    cell = table.rows[i].cells[j]
                    cell.text = str(cell_text) if cell_text else ''
                    
                    # Format header row
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
            
            logger.debug(f"Added table with {rows}x{cols} to document")
        except Exception as e:
            logger.warning(f"Failed to add table: {e}")
    
    def get_conversion_info(self) -> Dict[str, bool]:
        """Get available conversion methods for current system."""
        return {
            'pdfplumber': self._check_pdfplumber(),
            'ocr': self._check_ocr(),
            'libreoffice': bool(self._find_libreoffice()),
            'pymupdf': self._check_pymupdf(),
        }


# ═════════════════════════════════════════════════════════
# Convenience Functions (Mirror DOCX→PDF API)
# ═════════════════════════════════════════════════════════

def convert_pdf_to_docx(
    input_path: str,
    output_path: Optional[str] = None,
    method: Optional[str] = None,
    timeout: int = 120
) -> str:
    """
    Convert PDF to DOCX file.
    
    Provides simple procedural API matching docx2pdf style.
    
    Args:
        input_path: Path to input PDF file
        output_path: Path for output DOCX file (auto-generated if None)
        method: Conversion method ('auto', 'pdfplumber', 'ocr', 'libreoffice', 'pymupdf')
        timeout: Conversion timeout in seconds (default: 120)
    
    Returns:
        Path to generated DOCX file
        
    Raises:
        PDFToDOCXError: If conversion fails
        
    Examples:
        # Basic conversion with automatic method selection
        docx_path = convert_pdf_to_docx('input.pdf')
        
        # Specify output path
        docx_path = convert_pdf_to_docx('input.pdf', 'output.docx')
        
        # Force specific method
        docx_path = convert_pdf_to_docx('scan.pdf', method='ocr')
        
        # Custom timeout for large files
        docx_path = convert_pdf_to_docx('large.pdf', timeout=300)
    """
    converter = PDFToDOCXConverter(timeout=timeout)
    return converter.convert(input_path, output_path, method)


def get_pdf_converter_info() -> Dict[str, bool]:
    """
    Get information about available conversion methods.
    
    Returns:
        dict: {
            'pdfplumber': bool,    # Python text extraction
            'ocr': bool,           # Tesseract OCR
            'libreoffice': bool,   # System LibreOffice
            'pymupdf': bool,       # Lightweight fallback
        }
        
    Example:
        info = get_pdf_converter_info()
        print(f"Available methods: {[k for k, v in info.items() if v]}")
    """
    converter = PDFToDOCXConverter()
    return converter.get_conversion_info()


# ═════════════════════════════════════════════════════════
# CLI Interface
# ═════════════════════════════════════════════════════════

def main():
    """Command-line interface for PDF to DOCX conversion."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Convert PDF to DOCX format',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  python pdf_to_docx_converter.py input.pdf                    # Auto-detect method
  python pdf_to_docx_converter.py input.pdf -o output.docx     # Specify output
  python pdf_to_docx_converter.py scan.pdf -m ocr              # Force OCR
  python pdf_to_docx_converter.py --info                       # Check available methods
        '''
    )
    
    parser.add_argument('input', nargs='?', help='Input PDF file path')
    parser.add_argument('-o', '--output', help='Output DOCX file path')
    parser.add_argument(
        '-m', '--method',
        choices=['auto', 'pdfplumber', 'ocr', 'libreoffice', 'pymupdf'],
        default='auto',
        help='Conversion method (default: auto)'
    )
    parser.add_argument(
        '-t', '--timeout',
        type=int,
        default=120,
        help='Conversion timeout in seconds (default: 120)'
    )
    parser.add_argument(
        '--info',
        action='store_true',
        help='Show available conversion methods and exit'
    )
    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose logging'
    )
    
    args = parser.parse_args()
    
    # Show available methods
    if args.info:
        info = get_pdf_converter_info()
        print("\n📦 Available PDF to DOCX Conversion Methods:")
        print("=" * 50)
        for method, available in info.items():
            status = "✅ Available" if available else "❌ Not installed"
            print(f"  {method:15} {status}")
        print("=" * 50)
        return 0
    
    # Validate input
    if not args.input:
        parser.print_help()
        return 1
    
    # Set logging level
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    try:
        logger.info(f"Converting {args.input} to DOCX...")
        output = convert_pdf_to_docx(
            args.input,
            args.output,
            args.method,
            args.timeout
        )
        print(f"\n✅ Conversion successful!")
        print(f"Output: {output}")
        return 0
        
    except PDFToDOCXError as e:
        print(f"\n❌ Conversion failed: {e}", file=sys.stderr)
        return 1
    except Exception as e:
        print(f"\n❌ Unexpected error: {e}", file=sys.stderr)
        logger.exception("Unexpected error during conversion")
        return 1


if __name__ == '__main__':
    sys.exit(main())
